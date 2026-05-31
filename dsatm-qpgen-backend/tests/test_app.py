import os
import tempfile
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from fastapi.testclient import TestClient

os.environ["DATABASE_URL"] = "sqlite://"

test_workspace_root = Path(tempfile.gettempdir()) / "qpgen_backend_test"
os.environ["STORAGE_ROOT"] = str(test_workspace_root / "test-storage")

test_workspace_root.mkdir(parents=True, exist_ok=True)
Path(os.environ["STORAGE_ROOT"]).mkdir(parents=True, exist_ok=True)

from app.main import app
from app.database import SessionLocal
from app.models import Question, Subject, TeacherSubject, User

def login(client: TestClient, email: str, password: str) -> str:
    response = client.post("/api/v1/auth/login", json={"email": email, "password": password})
    assert response.status_code == 200
    return response.json()["access_token"]


def _question_table(document: DocxDocument):
    return next(
        table
        for table in document.tables
        if table.rows
        and table.rows[0].cells
        and table.rows[0].cells[0].text.strip().replace("\n", " ") == "Q No"
    )


def test_teacher_to_hod_review_flow() -> None:
    with TestClient(app) as client:
        teacher_token = login(client, "teacher@dsatm.edu", "Teacher@123")

        generate_response = client.post(
            "/api/v1/papers/generate",
            headers={"Authorization": f"Bearer {teacher_token}"},
            json={
                "subject_id": 1,
                "title": "IAT-1 Question Paper",
                "exam_type": "IAT-1",
                "semester": "5",
                "batch": "2022-26",
                "max_marks": 30,
                "duration_minutes": 90,
                "teaching_department": "AIML",
                "prompt": "Generate a balanced paper for CO1-CO3 with a mix of L1-L4.",
                "rbt_levels": ["L1", "L2", "L3", "L4"],
                "module_numbers": [1, 2, 3, 4],
            },
        )
        assert generate_response.status_code == 200
        paper = generate_response.json()

        submit_response = client.post(
            f"/api/v1/papers/{paper['id']}/submit",
            headers={"Authorization": f"Bearer {teacher_token}"},
        )
        assert submit_response.status_code == 200
        assert submit_response.json()["status"] == "pending_review"

        hod_token = login(client, "hod@dsatm.edu", "Hod@123")
        review_response = client.post(
            f"/api/v1/reviews/{paper['id']}/action",
            headers={"Authorization": f"Bearer {hod_token}"},
            json={"decision": "approved", "comments": "Balanced and ready for exam cell."},
        )
        assert review_response.status_code == 200
        assert review_response.json()["status"] == "approved"

        download_response = client.get(
            f"/api/v1/papers/{paper['id']}/download",
            headers={"Authorization": f"Bearer {hod_token}"},
        )
        assert download_response.status_code == 200
        exported = DocxDocument(BytesIO(download_response.content))
        paragraph_text = "\n".join(paragraph.text for paragraph in exported.paragraphs)
        table_text = "\n".join(
            cell.text for table in exported.tables for row in table.rows for cell in row.cells
        )
        assert "Dayananda Sagar Academy of Technology & Management" in table_text
        assert "USN:" in table_text
        assert "Department of Artificial Intelligence and Machine Learning" in paragraph_text or table_text
        assert "Percentage of CO Coverage" in paragraph_text
        assert "Percentage of Syllabus coverage" in paragraph_text
        assert "COs" in table_text
        assert "RBTL" in table_text


def test_paper_update_persists_overrides_and_custom_images() -> None:
    custom_image = (
        "data:image/png;base64,"
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9l9r8AAAAASUVORK5CYII="
    )

    with TestClient(app) as client:
        teacher_token = login(client, "teacher@dsatm.edu", "Teacher@123")
        headers = {"Authorization": f"Bearer {teacher_token}"}

        generate_response = client.post(
            "/api/v1/papers/generate",
            headers=headers,
            json={
                "subject_id": 1,
                "title": "Editable Paper",
                "exam_type": "IAT-1",
                "semester": "5",
                "batch": "2022-26",
                "max_marks": 30,
                "duration_minutes": 90,
                "teaching_department": "AIML",
                "prompt": "Generate an editable question paper.",
                "rbt_levels": ["L1", "L2", "L3", "L4"],
                "module_numbers": [1, 2, 3, 4],
            },
        )
        assert generate_response.status_code == 200
        paper = generate_response.json()
        first_question = paper["questions"][0]

        updated_text = "Design a precise AI workflow for solving the given problem with justification."
        update_response = client.put(
            f"/api/v1/papers/{paper['id']}",
            headers=headers,
            json={
                "question_text_overrides": {
                    str(first_question["id"]): updated_text,
                },
                "question_updates": [
                    {
                        "id": first_question["id"],
                        "text": updated_text,
                        "course_outcome": "CO5",
                        "bloom_level": "L6",
                        "module_number": 4,
                        "attached_images": [
                            {
                                "is_custom": True,
                                "file_name": "custom.png",
                                "caption": "Custom figure",
                                "data_url": custom_image,
                            }
                        ],
                    }
                ],
            },
        )
        assert update_response.status_code == 200, update_response.text
        updated_question = update_response.json()["questions"][0]
        assert updated_question["text"] == updated_text
        assert updated_question["course_outcome"] == "CO5"
        assert updated_question["bloom_level"] == "L6"
        assert updated_question["module_number"] == 4
        assert updated_question["attached_images"][0]["data_url"] == custom_image

        preview_response = client.get(
            f"/api/v1/papers/{paper['id']}/preview",
            headers=headers,
        )
        assert preview_response.status_code == 200
        preview_question = preview_response.json()["questions"][0]
        assert preview_question["text"] == updated_text
        assert preview_question["attached_images"][0]["caption"] == "Custom figure"

        download_response = client.get(
            f"/api/v1/papers/{paper['id']}/download",
            headers=headers,
        )
        assert download_response.status_code == 200
        exported = DocxDocument(BytesIO(download_response.content))
        assert len(exported.inline_shapes) >= 1


def test_end_sem_download_places_or_between_alternative_questions() -> None:
    with TestClient(app) as client:
        teacher_token = login(client, "teacher@dsatm.edu", "Teacher@123")

        generate_response = client.post(
            "/api/v1/ai/generate-paper",
            headers={"Authorization": f"Bearer {teacher_token}"},
            json={
                "subject_id": 3,
                "title": "End-Sem NLP Paper",
                "exam_type": "End-Sem",
                "semester": "6",
                "batch": "2022-26",
                "max_marks": 100,
                "duration_minutes": 180,
                "exam_date": "2026-04-30",
                "teaching_department": "AIML",
                "prompt": "Generate a module-balanced end-sem paper for NLP.",
                "rbt_levels": ["L1", "L2", "L3", "L4", "L5", "L6"],
                "module_numbers": [1, 2, 3, 4, 5],
                "difficulty_distribution": {"easy": 30, "medium": 40, "hard": 30},
                "co_targets": {"CO1": 20, "CO2": 20, "CO3": 20, "CO4": 20, "CO5": 20},
            },
        )
        assert generate_response.status_code == 200
        paper = generate_response.json()
        assert len(paper["questions"]) == 26
        assert len({item["text"] for item in paper["questions"]}) == 26
        assert all(not item.get("validation_errors") for item in paper["questions"])
        assert all(not item.get("validation_warnings") for item in paper["questions"])

        download_response = client.get(
            f"/api/v1/papers/{paper['id']}/download",
            headers={"Authorization": f"Bearer {teacher_token}"},
        )
        assert download_response.status_code == 200
        exported = DocxDocument(BytesIO(download_response.content))
        question_table = _question_table(exported)

        labels = [row.cells[0].text.strip() for row in question_table.rows]
        question_rows = {
            row.cells[0].text.strip(): row.cells[1].text.strip()
            for row in question_table.rows
            if row.cells[0].text.strip() and row.cells[0].text.strip()[0].isdigit()
        }
        assert len(question_rows) == 26
        assert all(question_rows.values())

        expected_breaks = [
            ("1(c)", "2(a)"),
            ("3(c)", "4(a)"),
            ("5(c)", "6(a)"),
            ("7(b)", "8(a)"),
            ("9(b)", "10(a)"),
        ]
        for previous_label, next_label in expected_breaks:
            previous_index = labels.index(previous_label)
            assert labels[previous_index + 1] == "OR"
            assert labels[previous_index + 2] == next_label


def test_manual_generation_respects_selected_question_order() -> None:
    with TestClient(app) as client:
        teacher_token = login(client, "teacher@dsatm.edu", "Teacher@123")

        questions_response = client.get(
            "/api/v1/questions?subject_id=1",
            headers={"Authorization": f"Bearer {teacher_token}"},
        )
        assert questions_response.status_code == 200
        manual_ids = [item["id"] for item in questions_response.json()[:20]]
        assert len(manual_ids) == 20

        generate_response = client.post(
            "/api/v1/ai/generate-paper",
            headers={"Authorization": f"Bearer {teacher_token}"},
            json={
                "subject_id": 1,
                "title": "Manual ML Paper",
                "exam_type": "IAT-2",
                "semester": "5",
                "batch": "2022-26",
                "max_marks": 50,
                "duration_minutes": 90,
                "teaching_department": "AIML",
                "prompt": "Use the manually curated question list.",
                "rbt_levels": ["L1", "L2", "L3", "L4"],
                "module_numbers": [1, 2, 3, 4, 5],
                "manual_question_ids": manual_ids,
            },
        )
        assert generate_response.status_code == 200
        paper = generate_response.json()
        assert [item["question_id"] for item in paper["questions"]] == manual_ids
        assert paper["questions"][0]["section_label"] == "1(a)"
        assert paper["questions"][1]["section_label"] == "1(b)"
        assert paper["questions"][2]["section_label"] == "2(a)"

        download_response = client.get(
            f"/api/v1/papers/{paper['id']}/download",
            headers={"Authorization": f"Bearer {teacher_token}"},
        )
        assert download_response.status_code == 200
        exported = DocxDocument(BytesIO(download_response.content))
        question_table = _question_table(exported)
        labels = [row.cells[0].text.strip() for row in question_table.rows]
        first_or_index = labels.index("1(b)") + 1
        assert labels[first_or_index] == "OR"
        assert labels[first_or_index + 1] == "2(a)"


def test_ai_paper_sanitizes_malformed_bank_questions() -> None:
    malformed_by_module = {
        1: [
            "Explain BAD402- Artificial Intelligence with suitable examples.",
            "Define the agent perceived the given percept at the given time with suitable examples.",
            "Solve give a Manhattan distance of h2 = 3+1 + 2 + 2+ 2 + 3+ 3 + 2 = 18 in the context of BAD402- Artificial Intelligence.",
            "Explain 5. Knowledge-based systems: The key to power? (1969-1979) with suitable examples.",
            "Explain the characteristics of intelligent agents with suitable examples.",
            "Apply heuristic search to solve a representative AI problem.",
        ],
        2: [
            "Define the workflow of Artificial Intelligence 1 with a neat diagram and suitable example.",
            "Solve 1. PROBLEM-SOLVING AGENTS Module 2 PROBLEM-SOLVING in the context of BAD402- Artificial Intelligence.",
            "Define 1. PROBLEM-SOLVING AGENTS Module 2 PROBLEM-SOLVING with suitable examples.",
            "Explain logical reasoning and its role in reasoning and knowledge representation.",
            "Analyze reasoning using predicate logic with suitable examples.",
            "Differentiate between semantic networks and frames.",
        ],
        3: [
            "Explain learning paradigms and its role in machine learning basics.",
            "Solve Module 3 in the context of FIRST-ORDER LOGIC . REPRESENTATION REVISITED.",
            "Compare learning paradigms and supervised learning for machine learning basics tasks.",
            "Analyze design trade-offs in model evaluation for scalable introduction to artificial intelligence systems.",
            "Explain supervised learning with suitable examples.",
            "Assess the role of model evaluation in machine learning.",
        ],
        4: [
            "Design an end-to-end introduction to artificial intelligence solution using text processing, tokenization, and sentiment analysis.",
            "Evaluate the limitations of sentiment analysis and propose suitable improvements.",
            "Solve The effect of heuristic accuracy on performance to solve a suitable problem and show the major steps involved.",
            "Define FIRST-ORDER LOGIC . REPRESENTATION REVISITED with suitable examples.",
            "Explain tokenization and its role in text processing.",
            "Analyze the challenges of sentiment analysis in natural language processing.",
            "Discuss parts-of-speech tagging with suitable examples.",
        ],
        5: [
            "Design an end-to-end introduction to artificial intelligence solution using image processing, feature detection, and object recognition.",
            "Explain image processing and its role in computer vision fundamentals.",
            "Compare image processing and feature detection for computer vision fundamentals tasks.",
            "Justify how image classification improves reliability or accuracy in computer vision fundamentals.",
            "Explain feature detection with suitable examples.",
            "Assess the limitations of object recognition in computer vision.",
            "Discuss object recognition methods used in computer vision systems.",
        ],
    }

    with TestClient(app) as client:
        with SessionLocal() as db:
            teacher = db.query(User).filter(User.email == "teacher@dsatm.edu").one()
            existing_count = db.query(Subject).filter(Subject.code.like("BAD402-QA-SAFE%")).count()
            subject = Subject(
                dept_id=int(teacher.dept_id or 1),
                code=f"BAD402-QA-SAFE-{existing_count + 1}",
                name="Introduction to Artificial Intelligence",
                semester=5,
                credits=4,
                max_marks=50,
            )
            db.add(subject)
            db.flush()
            db.add(TeacherSubject(teacher_id=teacher.id, subject_id=subject.id))

            question_seed = 1
            for module_number, questions in malformed_by_module.items():
                for text in questions:
                    db.add(
                        Question(
                            subject_id=subject.id,
                            teacher_id=teacher.id,
                            text=text,
                            marks=5 if question_seed % 2 else 6,
                            course_outcome=f"CO{((question_seed - 1) % 5) + 1}",
                            bloom_level=["L1", "L2", "L3", "L4"][question_seed % 4],
                            difficulty="medium",
                            module_number=module_number,
                            tags=["seed", "malformed-bank"],
                            is_verified=True,
                        )
                    )
                    question_seed += 1
            db.commit()
            subject_id = subject.id

        teacher_token = login(client, "teacher@dsatm.edu", "Teacher@123")

        generate_response = client.post(
            "/api/v1/ai/generate-paper",
            headers={"Authorization": f"Bearer {teacher_token}"},
            json={
                "subject_id": subject_id,
                "title": "AI Quality Gate Paper",
                "exam_type": "IAT-1",
                "semester": "5",
                "batch": "2022-26",
                "max_marks": 50,
                "duration_minutes": 90,
                "teaching_department": "AIML",
                "prompt": "Generate a clean, professional AI paper using only publishable questions.",
                "rbt_levels": ["L1", "L2", "L3", "L4"],
                "module_numbers": [1, 2, 3, 4, 5],
                "use_notes": False,
                "use_question_bank": True,
                "use_previous_papers": False,
                "use_syllabus": False,
            },
        )
        assert generate_response.status_code == 200, generate_response.text
        paper = generate_response.json()
        assert paper["questions"]
        assert all(not item.get("validation_errors") for item in paper["questions"])
        assert all(not item.get("validation_warnings") for item in paper["questions"])

        lowered_questions = [item["text"].lower() for item in paper["questions"]]
        forbidden_fragments = [
            "bad402",
            "given percept at the given time",
            "h2 =",
            "the key to power",
            "representation revisited",
            "problem-solving agents module",
        ]
        for fragment in forbidden_fragments:
            assert all(fragment not in question for question in lowered_questions)


def test_end_sem_intro_ai_generation_avoids_corrupted_bank_rows_and_wrong_syllabus_topics() -> None:
    malformed_rows = [
        (1, 6, "L1", "CO1", "Define the workflow of Artificial Intelligence Introduction to Artificial Intelligence • Homo Sapiens with a neat diagram and suitable example."),
        (1, 8, "L3", "CO3", "Solve the workflow of • The MAKE-ACTION-QUERY generates a sentence to ask which action should be done at the current time with a neat diagram and suitable example."),
        (2, 5, "L2", "CO1", "Explain the application of The left legs of Richard 4 in queries are employs a ima."),
        (2, 8, "L1", "CO5", "Define the workflow of computer passes the test if a human interrogator, after posing some written questions, cannot tell whether the written r with a neat diagram and suitable example."),
        (2, 7, "L3", "CO2", "Solve the workflow of Q is true in m with a neat diagram and suitable example."),
        (3, 5, "L1", "CO3", "Define the application of INFORMED (HEURISTIC) SEARCH STRATEGIES Informed search strategy—one that uses problem-specific knowledge beyond the definition of itself."),
        (4, 6, "L1", "CO4", "Define FIRST-ORDER LOGIC REPRESENTATION REVISITED with suitable examples."),
        (5, 8, "L3", "CO3", "Solve The effect of heuristic accuracy on performance to solve a suitable problem and show the major steps involved."),
    ]

    with TestClient(app) as client:
        with SessionLocal() as db:
            teacher = db.query(User).filter(User.email == "teacher@dsatm.edu").one()
            subject = db.query(Subject).filter(Subject.code == "21AI11").one()

            for module_number, marks, bloom_level, co, text in malformed_rows:
                db.add(
                    Question(
                        subject_id=subject.id,
                        teacher_id=teacher.id,
                        text=text,
                        marks=marks,
                        course_outcome=co,
                        bloom_level=bloom_level,
                        difficulty="medium",
                        module_number=module_number,
                        tags=["rag-generated", "corrupted-sample"],
                        is_verified=False,
                    )
                )
            db.commit()
            subject_id = subject.id

        teacher_token = login(client, "teacher@dsatm.edu", "Teacher@123")

        generate_response = client.post(
            "/api/v1/ai/generate-paper",
            headers={"Authorization": f"Bearer {teacher_token}"},
            json={
                "subject_id": subject_id,
                "title": "End-Sem AI Quality Gate Paper",
                "exam_type": "End-Sem",
                "semester": "4",
                "batch": "2022-26",
                "max_marks": 100,
                "duration_minutes": 180,
                "exam_date": "2026-05-06",
                "teaching_department": "Artificial Intelligence & Machine Learning",
                "prompt": "Generate a production-grade end-sem paper aligned to the official Introduction to Artificial Intelligence syllabus.",
                "rbt_levels": ["L1", "L2", "L3", "L4", "L5", "L6"],
                "module_numbers": [1, 2, 3, 4, 5],
                "co_targets": {"CO1": 20, "CO2": 20, "CO3": 20, "CO4": 20, "CO5": 20},
                "use_notes": True,
                "use_question_bank": True,
                "use_previous_papers": False,
                "use_syllabus": True,
            },
        )
        assert generate_response.status_code == 200, generate_response.text
        paper = generate_response.json()
        assert len(paper["questions"]) == 26
        assert all(not item.get("validation_errors") for item in paper["questions"])
        assert all(not item.get("validation_warnings") for item in paper["questions"])

        lowered_questions = [item["text"].lower() for item in paper["questions"]]
        forbidden_fragments = [
            "homo sapiens",
            "make-action-query",
            "the left legs of richard",
            "cannot tell whether the written",
            "q is true in m",
            "representation revisited",
        ]
        for fragment in forbidden_fragments:
            assert all(fragment not in question for question in lowered_questions)

        wrong_syllabus_topics = [
            "sentiment analysis",
            "object recognition",
            "supervised learning",
            "image processing",
        ]
        for fragment in wrong_syllabus_topics:
            assert all(fragment not in question for question in lowered_questions)
