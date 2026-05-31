from __future__ import annotations

import base64
import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .academic.templates import (
    DEFAULT_META_FIELDS,
    DSATM_TEMPLATE,
    HeaderLayout,
    InstitutionalTemplate,
    MetaField,
    SectionStyle,
    resolve_template_asset_path,
)


@dataclass
class PaperConfig:
    department: str
    subject: str
    subject_code: str
    semester: str
    max_marks: int
    duration: str
    date: str
    batch: str
    teaching_department: str
    exam_type: str
    modules: list[int]
    rbt_levels: list[str]
    co_targets: list[str]
    year: str = "2026"
    instructions: str = "Instruction: Answer the following questions"
    college_name: str = "Dayananda Sagar Academy of Technology & Management"
    affiliation: str = "(Autonomous Institute under VTU)"
    program_line: str = "6 Programs Accredited by NBA (CSE, ISE, ECE, EEE, MECH, CV)"
    co_descriptions: dict[str, str] = field(default_factory=dict)
    co_percentages: dict[str, int] = field(default_factory=dict)
    module_percentages: dict[str, int] = field(default_factory=dict)
    left_seal_label: str = "DSATM"
    right_seal_label: str = "IQAC"
    template_note: str | None = None
    template_family: str = "dsatm"
    template_config: dict[str, Any] = field(default_factory=dict)


def build_question_blueprint(max_marks: int) -> list[dict[str, Any]]:
    blueprint: list[dict[str, Any]] = []
    if max_marks <= 50:
        patterns = [(5, 5)] * 4 + [(4, 6)] * 6
        for question_number, (part_a, part_b) in enumerate(patterns, start=1):
            module_number = ((question_number - 1) // 2) + 1
            blueprint.extend(
                [
                    {
                        "question_number": question_number,
                        "subpart": "a",
                        "label": format_question_label(question_number, "a"),
                        "marks": part_a,
                        "module_number": module_number,
                    },
                    {
                        "question_number": question_number,
                        "subpart": "b",
                        "label": format_question_label(question_number, "b"),
                        "marks": part_b,
                        "module_number": module_number,
                    },
                ]
            )
        return blueprint

    hundred_mark_modules = {
        1: [(1, [("a", 6), ("b", 6), ("c", 8)]), (2, [("a", 6), ("b", 6), ("c", 8)])],
        2: [(3, [("a", 5), ("b", 8), ("c", 7)]), (4, [("a", 5), ("b", 8), ("c", 7)])],
        3: [(5, [("a", 5), ("b", 8), ("c", 7)]), (6, [("a", 5), ("b", 8), ("c", 7)])],
        4: [(7, [("a", 10), ("b", 10)]), (8, [("a", 10), ("b", 10)])],
        5: [(9, [("a", 10), ("b", 10)]), (10, [("a", 10), ("b", 10)])],
    }
    for module_number, question_sets in hundred_mark_modules.items():
        for question_number, parts in question_sets:
            for subpart, marks in parts:
                blueprint.append(
                    {
                        "question_number": question_number,
                        "subpart": subpart,
                        "label": format_question_label(question_number, subpart),
                        "marks": marks,
                        "module_number": module_number,
                    }
                )
    return blueprint


def format_question_label(question_number: int, subpart: str) -> str:
    return f"{question_number}({subpart})"


def normalize_question_label(label: Any) -> str:
    text = str(label or "").strip()
    if len(text) >= 2 and text[-1:].isalpha() and text[:-1].isdigit():
        return format_question_label(int(text[:-1]), text[-1:])
    return text


def build_question_rows(
    max_marks: int, questions: list[dict[str, Any]]
) -> list[dict[str, Any]]:
    blueprint = build_question_blueprint(max_marks)
    padded_questions = questions[: len(blueprint)] + [
        {} for _ in range(max(0, len(blueprint) - len(questions[: len(blueprint)])))
    ]
    rows: list[dict[str, Any]] = []
    current_module: int | None = None

    for slot, question in zip(blueprint, padded_questions):
        if max_marks > 50 and slot["module_number"] != current_module:
            current_module = int(slot["module_number"])
            rows.append({"type": "module", "title": f"Module - {current_module}"})

        if slot["subpart"] == "a" and slot["question_number"] % 2 == 0:
            rows.append({"type": "or"})

        rows.append(
            {
                "type": "question",
                "qno": normalize_question_label(
                    question.get("section_label") or slot["label"]
                ),
                "text": str(question.get("text", "")),
                "marks": int(question.get("marks", slot["marks"]) or slot["marks"]),
                "co": str(question.get("course_outcome", "")),
                "rbtl": str(question.get("bloom_level", "")),
                "attached_images": list(question.get("attached_images") or []),
            }
        )

    return rows


def _strip_diagram_placeholder(text: str) -> tuple[str, list[str]]:
    import re

    placeholders = re.findall(r"\[DIAGRAM:\s*(.*?)\]", str(text or ""), flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*\[DIAGRAM:\s*.*?\]\s*", " ", str(text or ""), flags=re.IGNORECASE)
    return " ".join(cleaned.split()).strip(), [path.strip() for path in placeholders if path.strip()]


def _decode_attachment_data_url(data_url: str) -> io.BytesIO | None:
    try:
        if not data_url.startswith("data:image/") or ";base64," not in data_url:
            return None
        _, encoded = data_url.split(",", 1)
        return io.BytesIO(base64.b64decode(encoded))
    except Exception:
        return None


def _resolve_question_images(item: dict[str, Any]) -> list[Any]:
    import os

    text, placeholder_paths = _strip_diagram_placeholder(str(item.get("text", "")))
    item["text"] = text
    resolved: list[tuple[str, Any]] = []
    for image in list(item.get("attached_images") or []):
        path = str(image.get("image_path") or "").strip()
        if path and os.path.exists(path):
            resolved.append((f"path:{path}", path))
            continue
        data_url = str(image.get("data_url") or "").strip()
        if data_url:
            image_stream = _decode_attachment_data_url(data_url)
            if image_stream is not None:
                resolved.append((f"data:{hash(data_url)}", image_stream))
    for path in placeholder_paths:
        if path and os.path.exists(path):
            resolved.append((f"path:{path}", path))
    deduped: list[Any] = []
    seen: set[str] = set()
    for key, asset in resolved:
        if key in seen:
            continue
        seen.add(key)
        deduped.append(asset)
    return deduped[:2]


class DSATMQuestionPaperGenerator:
    def generate(
        self, config: PaperConfig, questions: list[dict[str, Any]]
    ) -> DocumentType:
        self.template = self._coerce_template(config.template_config)
        document = Document()
        self._set_page_layout(document, self.template)
        self._add_header(document, config, self.template)
        if self.template.show_usn_row:
            self._add_usn_row(document, self.template.usn_box_count)
        if self.template.show_department_heading:
            self._add_department_heading(document, config, self.template)
        self._add_exam_title(document, config, self.template)
        self._add_meta_table(document, config, self.template)
        self._add_instruction(document, config, self.template)
        if self.template.question_table_style == SectionStyle.NUMBERED_LIST:
            self._add_questions_list(document, config, questions, self.template)
        else:
            self._add_questions_table(document, config, questions, self.template)
        if self.template.show_co_descriptions:
            self._add_course_outcomes_table(document, config, self.template)
        if self.template.show_co_coverage_table or self.template.show_module_coverage_table:
            self._add_coverage_page(document, config, self.template)
        return document

    def save(self, document: DocumentType, output_path: Path) -> Path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))
        return output_path

    def _coerce_template(self, data: dict[str, Any] | None) -> InstitutionalTemplate:
        if not data:
            return DSATM_TEMPLATE
        try:
            return InstitutionalTemplate.from_dict(dict(data))
        except Exception:
            return DSATM_TEMPLATE

    def _hex_to_rgb(self, color: str | None) -> RGBColor | None:
        value = str(color or "").strip().lstrip("#")
        if len(value) != 6:
            return None
        try:
            return RGBColor.from_string(value.upper())
        except ValueError:
            return None

    def _meta_value(self, config: PaperConfig, key: str) -> str:
        mapping = {
            "subject_name": config.subject,
            "subject_code": config.subject_code,
            "semester": config.semester,
            "max_marks": str(config.max_marks),
            "batch": config.batch,
            "duration": config.duration,
            "exam_date": config.date,
            "teaching_department": config.teaching_department,
            "department": config.department,
            "exam_type": config.exam_type,
            "rbt_levels_text": (
                "L1-Remember, L2-Understand, L3-Apply, "
                "L4-Analyze, L5-Evaluate, L6-Create"
            ),
        }
        return str(mapping.get(key, ""))

    def _set_page_layout(
        self, document: DocumentType, template: InstitutionalTemplate
    ) -> None:
        section = document.sections[0]
        section.page_width = Inches(template.page_width_inches)
        section.page_height = Inches(template.page_height_inches)
        section.top_margin = Inches(template.margin_top_inches)
        section.bottom_margin = Inches(template.margin_bottom_inches)
        section.left_margin = Inches(template.margin_left_inches)
        section.right_margin = Inches(template.margin_right_inches)
        section.start_type = WD_SECTION.NEW_PAGE

        normal = document.styles["Normal"]
        normal.font.name = template.font_family
        normal.font.size = Pt(template.base_font_size)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), template.font_family)

    def _style_run(
        self,
        run,
        *,
        size: int = 9,
        bold: bool = False,
        italic: bool = False,
        color: RGBColor | None = None,
    ) -> None:
        font_family = getattr(self, "template", DSATM_TEMPLATE).font_family
        run.font.name = font_family
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_family)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color

    def _set_table_borders(
        self,
        table,
        *,
        top: bool = True,
        left: bool = True,
        bottom: bool = True,
        right: bool = True,
        inside_h: bool = True,
        inside_v: bool = True,
        size: str = "10",
    ) -> None:
        borders = OxmlElement("w:tblBorders")
        mapping = {
            "top": top,
            "left": left,
            "bottom": bottom,
            "right": right,
            "insideH": inside_h,
            "insideV": inside_v,
        }
        for name, enabled in mapping.items():
            border = OxmlElement(f"w:{name}")
            border.set(qn("w:val"), "single" if enabled else "nil")
            border.set(qn("w:sz"), size)
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            borders.append(border)

        table_element = table._tbl
        table_pr = table_element.tblPr
        existing = table_pr.first_child_found_in("w:tblBorders")
        if existing is not None:
            table_pr.remove(existing)
        table_pr.append(borders)

    def _set_table_fixed_layout(self, table) -> None:
        table_pr = table._tbl.tblPr
        existing = table_pr.first_child_found_in("w:tblLayout")
        if existing is not None:
            table_pr.remove(existing)
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        table_pr.append(layout)

    def _set_cell_margins(
        self,
        cell,
        *,
        top: int = 45,
        left: int = 65,
        bottom: int = 45,
        right: int = 65,
    ) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        margins = tc_pr.first_child_found_in("w:tcMar")
        if margins is None:
            margins = OxmlElement("w:tcMar")
            tc_pr.append(margins)
        for side, value in {
            "top": top,
            "start": left,
            "bottom": bottom,
            "end": right,
        }.items():
            node = margins.find(qn(f"w:{side}"))
            if node is None:
                node = OxmlElement(f"w:{side}")
                margins.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    def _set_cell(
        self,
        cell,
        text: str,
        *,
        bold: bool = False,
        italic: bool = False,
        size: int = 9,
        align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
        vertical_align: WD_ALIGN_VERTICAL = WD_ALIGN_VERTICAL.CENTER,
    ) -> None:
        cell.text = ""
        self._set_cell_margins(cell)
        lines = str(text).splitlines() or [""]
        for index, line in enumerate(lines):
            paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
            paragraph.alignment = align
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(line)
            self._style_run(run, size=size, bold=bold, italic=italic)
        cell.vertical_alignment = vertical_align

    def _render_asset_or_label(
        self, cell, label: str, asset_path: str | None
    ) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        asset = resolve_template_asset_path(asset_path)
        if asset is not None:
            run = paragraph.add_run()
            run.add_picture(str(asset), width=Inches(0.62))
        elif label:
            top = paragraph.add_run(label)
            self._style_run(top, size=max(self.template.base_font_size, 10), bold=True)
            bottom = paragraph.add_run("\nSeal")
            self._style_run(bottom, size=max(self.template.base_font_size - 2, 7))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def _append_accreditation_line(
        self,
        paragraph,
        text: str,
        highlighted_parts: list[str],
        *,
        size: int,
        accent: RGBColor | None,
    ) -> None:
        remaining = text
        if not highlighted_parts:
            run = paragraph.add_run(text)
            self._style_run(run, size=size)
            return

        for highlighted in highlighted_parts:
            if not highlighted or highlighted not in remaining:
                continue
            before, after = remaining.split(highlighted, 1)
            if before:
                run = paragraph.add_run(before)
                self._style_run(run, size=size)
            focus = paragraph.add_run(highlighted)
            self._style_run(focus, size=size, color=accent)
            remaining = after

        if remaining:
            run = paragraph.add_run(remaining)
            self._style_run(run, size=size)

    def _add_header(
        self,
        document: DocumentType,
        config: PaperConfig,
        template: InstitutionalTemplate,
    ) -> None:
        accent = self._hex_to_rgb(template.accent_color)
        institution_name = template.institution_name or config.college_name
        affiliation = template.affiliation_text or config.affiliation

        if template.header_layout == HeaderLayout.BANNER:
            banner = resolve_template_asset_path(template.banner_image_path)
            if banner is not None:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run().add_picture(str(banner), width=Inches(7.1))
            title = document.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._style_run(title.add_run(institution_name), size=template.title_font_size, bold=True)
            if affiliation:
                title.add_run("\n")
                self._style_run(title.add_run(affiliation), size=max(template.base_font_size - 1, 8))
        elif template.header_layout == HeaderLayout.MINIMAL:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._style_run(paragraph.add_run(institution_name), size=template.title_font_size, bold=True)
            if affiliation:
                paragraph.add_run("\n")
                self._style_run(paragraph.add_run(affiliation), size=max(template.base_font_size - 1, 8))
        elif template.header_layout == HeaderLayout.SINGLE_LOGO_CENTER:
            logo_para = document.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_path = template.left_seal_path or template.right_seal_path
            logo_asset = resolve_template_asset_path(logo_path)
            if logo_asset is not None:
                logo_para.add_run().add_picture(str(logo_asset), width=Inches(0.72))
            elif template.left_seal_label or template.right_seal_label:
                self._style_run(
                    logo_para.add_run(template.left_seal_label or template.right_seal_label),
                    size=max(template.header_font_size, 10),
                    bold=True,
                )

            title = document.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._style_run(title.add_run(institution_name), size=template.title_font_size, bold=True)
            if affiliation:
                title.add_run("\n")
                self._style_run(title.add_run(affiliation), size=max(template.base_font_size - 1, 8))
            if template.accreditation_lines:
                notes = document.add_paragraph()
                notes.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for line in template.accreditation_lines:
                    self._append_accreditation_line(
                        notes,
                        line.text,
                        line.highlighted_parts,
                        size=max(template.base_font_size - 1, 8),
                        accent=accent,
                    )
                    notes.add_run("\n")
        else:
            table = document.add_table(rows=1, cols=4)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            self._set_table_borders(
                table,
                top=False,
                left=False,
                bottom=False,
                right=False,
                inside_h=False,
                inside_v=False,
            )
            self._set_table_fixed_layout(table)

            widths = [0.8, 4.0, 1.7, 0.8]
            for index, width in enumerate(widths):
                table.columns[index].width = Inches(width)

            left_cell, title_cell, approval_cell, right_cell = table.rows[0].cells
            self._render_asset_or_label(left_cell, template.left_seal_label or config.left_seal_label, template.left_seal_path)

            title = title_cell.paragraphs[0]
            title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self._style_run(title.add_run(institution_name), size=template.header_font_size, bold=True)
            if affiliation:
                title.add_run("\n")
                self._style_run(title.add_run(affiliation), size=max(template.base_font_size - 1, 8))

            approval = approval_cell.paragraphs[0]
            approval.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lines = template.accreditation_lines or []
            for line in lines:
                self._append_accreditation_line(
                    approval,
                    line.text,
                    line.highlighted_parts,
                    size=max(template.base_font_size - 1, 7),
                    accent=accent,
                )
                approval.add_run("\n")

            self._render_asset_or_label(right_cell, template.right_seal_label or config.right_seal_label, template.right_seal_path)

        divider = document.add_paragraph()
        divider.paragraph_format.space_before = Pt(2)
        divider.paragraph_format.space_after = Pt(2)
        divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = divider.add_run("_" * 116)
        self._style_run(run, size=max(template.base_font_size - 2, 7))

    def _add_usn_row(self, document: DocumentType, box_count: int) -> None:
        total_columns = max(2, box_count + 1)
        table = document.add_table(rows=1, cols=total_columns)
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        table.autofit = False
        self._set_table_borders(table, top=False, left=False, bottom=False, right=False)
        self._set_table_fixed_layout(table)

        label_cell = table.rows[0].cells[0]
        label_cell.width = Inches(0.55)
        self._set_cell(
            label_cell,
            "USN:",
            size=max(self.template.base_font_size - 1, 8),
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        for index in range(1, total_columns):
            table.columns[index].width = Inches(0.31)
            self._set_cell(
                table.rows[0].cells[index],
                "",
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            self._set_table_borders(
                table,
                top=False,
                left=False,
                bottom=False,
                right=False,
                inside_h=False,
                inside_v=False,
            )
            cell = table.rows[0].cells[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "8")
                border.set(qn("w:color"), "000000")
                borders.append(border)
            tc_pr.append(borders)

    def _add_department_heading(
        self,
        document: DocumentType,
        config: PaperConfig,
        template: InstitutionalTemplate,
    ) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(4)
        prefix = template.department_prefix or "Department of"
        run = paragraph.add_run(f"{prefix} {config.department}".strip())
        self._style_run(run, size=max(template.title_font_size, 12), bold=True)

    def _add_exam_title(
        self,
        document: DocumentType,
        config: PaperConfig,
        template: InstitutionalTemplate,
    ) -> None:
        if not template.show_exam_title_box:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._style_run(
                paragraph.add_run(config.exam_type),
                size=max(template.header_font_size, 10),
                bold=True,
            )
            return
        table = document.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_borders(table, top=template.exam_title_bordered, left=template.exam_title_bordered, bottom=template.exam_title_bordered, right=template.exam_title_bordered)
        self._set_table_fixed_layout(table)
        table.columns[0].width = Inches(7.35)
        self._set_cell(
            table.rows[0].cells[0],
            config.exam_type,
            bold=True,
            size=max(template.header_font_size, 10),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    def _add_meta_table(
        self,
        document: DocumentType,
        config: PaperConfig,
        template: InstitutionalTemplate,
    ) -> None:
        left_fields = [
            field for field in (template.meta_fields or list(DEFAULT_META_FIELDS))
            if field.position == "left"
        ]
        right_fields = [
            field for field in (template.meta_fields or list(DEFAULT_META_FIELDS))
            if field.position == "right"
        ]
        full_fields = [
            field for field in (template.meta_fields or list(DEFAULT_META_FIELDS))
            if field.position == "full"
        ]
        row_count = max(len(left_fields), len(right_fields)) + len(full_fields)
        table = document.add_table(rows=max(row_count, 1), cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_borders(table)
        self._set_table_fixed_layout(table)
        widths = [1.5, 3.0, 1.5, 1.35]
        for index, width in enumerate(widths):
            table.columns[index].width = Inches(width)

        row_index = 0
        for index in range(max(len(left_fields), len(right_fields))):
            left = left_fields[index] if index < len(left_fields) else None
            right = right_fields[index] if index < len(right_fields) else None
            self._set_cell(table.rows[row_index].cells[0], left.label if left else "", bold=bool(left and left.bold_label))
            self._set_cell(table.rows[row_index].cells[1], self._meta_value(config, left.key) if left else "", bold=bool(left and left.bold_value))
            self._set_cell(table.rows[row_index].cells[2], right.label if right else "", bold=bool(right and right.bold_label))
            self._set_cell(table.rows[row_index].cells[3], self._meta_value(config, right.key) if right else "", bold=bool(right and right.bold_value))
            row_index += 1

        for field in full_fields:
            self._set_cell(table.rows[row_index].cells[0], field.label, bold=field.bold_label)
            merged = table.rows[row_index].cells[1].merge(table.rows[row_index].cells[3])
            self._set_cell(merged, self._meta_value(config, field.key), bold=field.bold_value)
            row_index += 1

    def _add_instruction(
        self,
        document: DocumentType,
        config: PaperConfig,
        template: InstitutionalTemplate,
    ) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(config.instructions or template.default_instructions)
        self._style_run(run, size=template.base_font_size, italic=True)

        if template.show_template_note and config.template_note:
            note = document.add_paragraph()
            note.paragraph_format.space_before = Pt(4)
            note.paragraph_format.space_after = Pt(0)
            note_label = note.add_run("Note:\n")
            self._style_run(note_label, size=template.base_font_size, bold=True)
            body = note.add_run(config.template_note)
            self._style_run(body, size=template.base_font_size, bold=True)

    def _add_questions_table(
        self,
        document: DocumentType,
        config: PaperConfig,
        questions: list[dict[str, Any]],
        template: InstitutionalTemplate,
    ) -> None:
        paper_rows = [
            row
            for row in build_question_rows(config.max_marks, questions)
            if (
                (row["type"] != "module" or template.show_module_headers)
                and (row["type"] != "or" or template.show_or_separators)
            )
        ]
        columns: list[tuple[str, str, float]] = []
        if template.show_qno_column:
            columns.append(("qno", "Q\nNo", 0.7))
        columns.append(("text", "Questions", 4.6))
        if template.show_marks_column:
            columns.append(("marks", "Marks", 0.8))
        if template.show_co_column:
            columns.append(("co", "COs", 0.7))
        if template.show_rbtl_column:
            columns.append(("rbtl", "RBTL", 0.7))

        if len(columns) == 1:
            columns.append(("text", "Questions", 6.8))

        table = document.add_table(rows=1 + len(paper_rows), cols=len(columns))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_borders(table)
        self._set_table_fixed_layout(table)
        widths = [item[2] for item in columns]
        total_width = sum(widths)
        if total_width < 7.7:
            text_index = next((index for index, item in enumerate(columns) if item[0] == "text"), 0)
            widths[text_index] += 7.7 - total_width
        for index, width in enumerate(widths):
            table.columns[index].width = Inches(width)

        for index, (_, header, _) in enumerate(columns):
            self._set_cell(
                table.rows[0].cells[index],
                header,
                bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                size=template.base_font_size,
            )

        row_index = 1
        for item in paper_rows:
            if item["type"] == "module":
                module_row = table.rows[row_index]
                merged = module_row.cells[0].merge(module_row.cells[len(columns) - 1])
                self._set_cell(
                    merged,
                    item["title"],
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    size=template.base_font_size,
                )
                row_index += 1
                continue

            if item["type"] == "or":
                or_row = table.rows[row_index]
                merged = or_row.cells[0].merge(or_row.cells[len(columns) - 1])
                self._set_cell(
                    merged,
                    "OR",
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    size=template.base_font_size,
                )
                row_index += 1
                continue

            current = table.rows[row_index]
            for cell_index, width in enumerate(widths):
                current.cells[cell_index].width = Inches(width)
            value_map = {
                "qno": item.get("qno", ""),
                "text": item.get("text", ""),
                "marks": str(item.get("marks", "")),
                "co": item.get("co", ""),
                "rbtl": item.get("rbtl", ""),
            }
            for column_index, (key, _, _) in enumerate(columns):
                cell = current.cells[column_index]
                val = str(value_map.get(key, ""))
                if key == "text":
                    item["text"] = val
                    image_paths = _resolve_question_images(item)
                    val = str(item.get("text", ""))
                    
                    self._set_cell(
                        cell,
                        val,
                        align=WD_ALIGN_PARAGRAPH.LEFT,
                        vertical_align=WD_ALIGN_VERTICAL.TOP,
                        size=template.base_font_size,
                    )
                    
                    for image_asset in image_paths:
                        p_img = cell.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.paragraph_format.space_before = Pt(6)
                        p_img.paragraph_format.space_after = Pt(6)
                        try:
                            p_img.add_run().add_picture(image_asset, width=Inches(3.5))
                        except Exception as e:
                            import logging
                            logging.getLogger("app.generator").error(f"Failed to add picture in DOCX cell: {e}")
                else:
                    self._set_cell(
                        cell,
                        val,
                        align=WD_ALIGN_PARAGRAPH.CENTER,
                        vertical_align=WD_ALIGN_VERTICAL.TOP,
                        size=template.base_font_size,
                    )
            current.height_rule = WD_ROW_HEIGHT_RULE.AUTO
            row_index += 1

    def _add_questions_list(
        self,
        document: DocumentType,
        config: PaperConfig,
        questions: list[dict[str, Any]],
        template: InstitutionalTemplate,
    ) -> None:
        paper_rows = [
            row
            for row in build_question_rows(config.max_marks, questions)
            if (
                (row["type"] != "module" or template.show_module_headers)
                and (row["type"] != "or" or template.show_or_separators)
            )
        ]
        for item in paper_rows:
            if item["type"] == "module":
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._style_run(paragraph.add_run(item["title"]), size=template.base_font_size, bold=True)
                continue
            if item["type"] == "or":
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._style_run(paragraph.add_run("OR"), size=template.base_font_size, bold=True)
                continue

            image_paths = _resolve_question_images(item)
            val = str(item.get("text", ""))

            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            prefix = f"{item['qno']} " if template.show_qno_column else ""
            self._style_run(paragraph.add_run(prefix), size=template.base_font_size, bold=True)
            self._style_run(paragraph.add_run(val), size=template.base_font_size)

            for image_asset in image_paths:
                p_img = document.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(6)
                p_img.paragraph_format.space_after = Pt(6)
                try:
                    p_img.add_run().add_picture(image_asset, width=Inches(3.5))
                except Exception as e:
                    import logging
                    logging.getLogger("app.generator").error(f"Failed to add picture in DOCX list: {e}")

            meta_bits: list[str] = []
            if template.show_marks_column:
                meta_bits.append(f"Marks: {item['marks']}")
            if template.show_co_column and item.get("co"):
                meta_bits.append(f"CO: {item['co']}")
            if template.show_rbtl_column and item.get("rbtl"):
                meta_bits.append(f"RBTL: {item['rbtl']}")
            if meta_bits:
                meta = document.add_paragraph()
                meta.paragraph_format.space_after = Pt(4)
                self._style_run(
                    meta.add_run(" | ".join(meta_bits)),
                    size=max(template.base_font_size - 1, 8),
                    italic=True,
                )

    def _add_course_outcomes_table(
        self,
        document: DocumentType,
        config: PaperConfig,
        template: InstitutionalTemplate,
    ) -> None:
        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(2)
        label = heading.add_run(
            "Course Outcomes (COs):  At the end of the Course, the Student will be able to:"
        )
        self._style_run(label, size=max(template.base_font_size - 1, 8), bold=True)

        co_keys = sorted(set(config.co_descriptions.keys()) or {f"CO{index}" for index in range(1, 6)})
        table = document.add_table(rows=max(len(co_keys), 1), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_borders(table)
        self._set_table_fixed_layout(table)
        table.columns[0].width = Inches(0.6)
        table.columns[1].width = Inches(6.7)
        for index, co_key in enumerate(co_keys, start=1):
            self._set_cell(
                table.rows[index - 1].cells[0],
                co_key,
                bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                size=template.base_font_size,
            )
            self._set_cell(
                table.rows[index - 1].cells[1],
                config.co_descriptions.get(co_key, ""),
                size=template.base_font_size,
            )

    def _add_coverage_page(
        self,
        document: DocumentType,
        config: PaperConfig,
        template: InstitutionalTemplate,
    ) -> None:
        document.add_page_break()
        if template.show_co_coverage_table:
            co_keys = sorted(set(config.co_percentages.keys()) or {f"CO{index}" for index in range(1, 6)})
            co_heading = document.add_paragraph()
            run = co_heading.add_run("Percentage of CO Coverage")
            self._style_run(run, size=template.base_font_size, bold=True)

            co_table = document.add_table(rows=2, cols=len(co_keys) + 1)
            co_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            co_table.autofit = False
            self._set_table_borders(co_table)
            self._set_table_fixed_layout(co_table)
            self._set_cell(co_table.rows[0].cells[0], "Course Outcomes", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=template.base_font_size)
            self._set_cell(co_table.rows[1].cells[0], "Percentage", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=template.base_font_size)
            for index, co_key in enumerate(co_keys, start=1):
                self._set_cell(co_table.rows[0].cells[index], co_key, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=template.base_font_size)
                self._set_cell(
                    co_table.rows[1].cells[index],
                    str(config.co_percentages.get(co_key, 0)),
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    size=template.base_font_size,
                )

        if template.show_module_coverage_table:
            module_heading = document.add_paragraph()
            module_heading.paragraph_format.space_before = Pt(12)
            run = module_heading.add_run("Percentage of Syllabus coverage")
            self._style_run(run, size=template.base_font_size, bold=True)

            module_keys = sorted(
                {
                    str(key)
                    for key in config.module_percentages.keys()
                }
                or {str(index) for index in range(1, 6)}
            )
            module_table = document.add_table(rows=2, cols=len(module_keys) + 1)
            module_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            module_table.autofit = False
            self._set_table_borders(module_table)
            self._set_table_fixed_layout(module_table)
            self._set_cell(module_table.rows[0].cells[0], "Modules Covered", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=template.base_font_size)
            self._set_cell(module_table.rows[1].cells[0], "Percentage", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=template.base_font_size)
            for index, module_key in enumerate(module_keys, start=1):
                self._set_cell(module_table.rows[0].cells[index], module_key, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=template.base_font_size)
                self._set_cell(
                    module_table.rows[1].cells[index],
                    str(config.module_percentages.get(module_key, config.module_percentages.get(int(module_key), 0))),
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    size=template.base_font_size,
                )


def generate_question_paper(
    config: PaperConfig, questions: list[dict[str, Any]], output_dir: Path
) -> Path:
    generator = DSATMQuestionPaperGenerator()
    document = generator.generate(config, questions)
    filename = (
        f"QP_{config.subject_code}_{config.exam_type}_{config.date.replace('-', '')}.docx"
    )
    return generator.save(document, output_dir / filename)


docx_generator = DSATMQuestionPaperGenerator()
